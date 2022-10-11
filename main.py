import docx.document
import docx.table
from scraping import get_word


def add_word(word_info: dict[str, str]):
    document: docx.document.Document = docx.Document('Glossary.docx')
    table: docx.table.Table = document.tables[0]
    table.add_row()
    cells: list[table.table.Cell] = table.rows[-1].cells

    cells[0].text = word_info['word']
    cells[0].add_paragraph(word_info['transcription'])

    cells[1].text = word_info['meaning']
    p = cells[1].add_paragraph()
    p.add_run(word_info['translation']).bold = True

    cells[2].text = word_info['example']

    cells[3].text = '\n'.join(word_info['synonyms'])

    cells[4].text = '\n'.join(word_info['antonyms'])

    document.save('Glossary.docx')


if __name__ == '__main__':
    for word in input().split(' '):
        try:
            info = get_word(word)
            if len(info) == 0:
                continue
            add_word(info)
        except BaseException:
            print("Problem with word", word)

    input()
